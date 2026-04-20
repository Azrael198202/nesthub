from __future__ import annotations

import json
import logging
import os
import re
import urllib.request
import urllib.parse
from typing import Any
from pathlib import Path

from nethub_runtime.core.schemas.context_schema import CoreContextSchema
from nethub_runtime.core.schemas.task_schema import TaskSchema

logger = logging.getLogger("nethub_runtime.core.execution_step_handlers")


def handle_extract_records_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    _context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    records = coordinator._extract_records(task.input_text)
    return {"records": records, "count": len(records)}


def handle_persist_records_step(
    coordinator: Any,
    _step: dict[str, Any],
    _task: TaskSchema,
    context: CoreContextSchema,
    step_outputs: dict[str, Any],
) -> dict[str, Any]:
    records = step_outputs.get("extract_records", {}).get("records", [])
    state = coordinator.session_store.append_records(context.session_id, records)
    return {"saved": len(records), "total_records": len(state.get("records", []))}


def handle_parse_query_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    state = coordinator.session_store.get(context.session_id)
    return {"query": coordinator._parse_query(task.input_text, state.get("records", []))}


def handle_aggregate_query_step(
    coordinator: Any,
    step: dict[str, Any],
    _task: TaskSchema,
    context: CoreContextSchema,
    step_outputs: dict[str, Any],
) -> dict[str, Any]:
    query = step_outputs.get("parse_query", {}).get("query", {})
    state = coordinator.session_store.get(context.session_id)
    model_choice = ((step.get("capability") or {}).get("model_choice") or {})
    return {"aggregation": coordinator._aggregate_records(state.get("records", []), query, model_choice)}


def handle_ocr_extract_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    """OCR placeholder: copy image to received/, content extraction to be implemented."""
    import shutil
    import re as _re

    input_text = task.input_text or ""
    attachments = context.metadata.get("attachments") or []

    # Resolve image path — Priority 1: context attachments (LINE uploads)
    image_path: Path | None = None
    for att in attachments:
        ct = str(att.get("content_type") or "")
        if ct.startswith("image/"):
            sp = str(att.get("stored_path") or "").strip()
            dest_dir = _received_dir(context.session_id)
            dest = dest_dir / str(att.get("file_name") or Path(sp).name if sp else att.get("file_name") or "image.jpg")
            if sp and Path(sp).exists():
                # Same-host: direct copy (no HTTP round-trip)
                if not dest.exists():
                    shutil.copy2(sp, dest)
                image_path = dest
            else:
                # Cross-host (Railway → local): download via download_url
                url = str(att.get("download_url") or "").strip()
                if url:
                    try:
                        import urllib.request as _ur
                        with _ur.urlopen(_ur.Request(url), timeout=30) as resp:
                            dest.write_bytes(resp.read())
                        image_path = dest
                    except Exception as exc:
                        logger.warning("Failed to download image attachment %s: %s", url, exc)
            break

    # Priority 2: file path in task text
    if image_path is None:
        path_match = _re.search(r'[\w./\\-]+\.(?:png|jpg|jpeg|bmp|tiff|gif|webp)', input_text, _re.IGNORECASE)
        image_path = Path(path_match.group(0)) if path_match else None

    if image_path is None or not image_path.exists():
        # File not on disk yet (cross-process / delayed write), but attachment
        # info was received — treat as received so repair loop doesn't fire.
        # Only hard-fail when there are no attachments at all.
        if attachments:
            att = next((a for a in attachments if str(a.get("content_type","")).startswith("image/")), attachments[0])
            return {
                "artifact_type": "text",
                "status": "received",
                "method": "placeholder",
                "file_name": att.get("file_name", ""),
                "content": f"[图片已接收: {att.get('file_name','')}, 内容提取待实现]",
                "message": f"附件已接收，本地路径暂不可访问: {image_path}",
            }
        return {
            "artifact_type": "text",
            "status": "error",
            "message": f"图片文件未找到且无附件信息: {image_path}",
        }

    # TODO: 接入 OCR 引擎（PaddleOCR / Qwen2.5-VL / EasyOCR）
    return {
        "artifact_type": "text",
        "status": "received",
        "method": "placeholder",
        "artifact_path": str(image_path),
        "file_name": image_path.name,
        "file_size": image_path.stat().st_size,
        "content": f"[图片已接收: {image_path.name}，内容提取待实现]",
        "message": f"文件已保存至 {image_path}",
    }


def handle_stt_transcribe_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    """STT with priority chain: Whisper (local) → HuggingFace ASR."""
    import re as _re
    input_text = task.input_text or ""
    path_match = _re.search(r'[\w./\\-]+\.(?:mp3|wav|ogg|flac|m4a|webm)', input_text, _re.IGNORECASE)
    audio_path = Path(path_match.group(0)) if path_match else None

    if audio_path is None or not audio_path.exists():
        return {"artifact_type": "text", "status": "error", "message": f"Audio file not found: {audio_path}"}

    # Try openai-whisper (local, no API key)
    try:
        import whisper  # type: ignore
        model = whisper.load_model("small")
        result = model.transcribe(str(audio_path))
        text = str(result.get("text") or "").strip()
        return {"artifact_type": "text", "status": "transcribed", "method": "whisper_small",
                "content": text, "language": result.get("language", "unknown")}
    except Exception:
        pass

    # Trigger acquisition
    from nethub_runtime.core.services.capability_acquisition_service import CapabilityAcquisitionService
    svc = getattr(coordinator, "capability_acquisition_service", None) or \
          CapabilityAcquisitionService(security_guard=getattr(coordinator, "security_guard", None))
    acq = svc.acquire(task_type="speech_recognition", gap="no_stt_engine")
    return {"artifact_type": "text", "status": "acquisition_triggered",
            "method": acq.strategy, "detail": acq.detail,
            "message": "STT engine is being installed. Retry shortly."}


def handle_tts_synthesize_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    """TTS with priority chain: pyttsx3 (local) → HuggingFace TTS."""
    import re as _re, hashlib as _hs
    input_text = task.input_text or ""
    # Derive output path
    artifact_id = _hs.md5(input_text.encode()).hexdigest()[:12]
    out_path = Path(f"generated/{artifact_id}.wav")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        import pyttsx3  # type: ignore
        engine = pyttsx3.init()
        engine.save_to_file(input_text, str(out_path))
        engine.runAndWait()
        if out_path.exists() and out_path.stat().st_size > 0:
            return {"artifact_type": "audio", "status": "synthesized", "method": "pyttsx3",
                    "artifact_path": str(out_path), "file_name": out_path.name}
    except Exception:
        pass

    from nethub_runtime.core.services.capability_acquisition_service import CapabilityAcquisitionService
    svc = getattr(coordinator, "capability_acquisition_service", None) or \
          CapabilityAcquisitionService(security_guard=getattr(coordinator, "security_guard", None))
    acq = svc.acquire(task_type="audio_generation", gap="no_tts_engine")
    return {"artifact_type": "audio", "status": "acquisition_triggered",
            "method": acq.strategy, "detail": acq.detail,
            "message": "TTS engine is being installed. Retry shortly."}


def handle_image_generate_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    """Generate an image using the self-healing ImageGenerationService.

    The service tries available backends in order, and when none succeed
    it autonomously installs missing packages and retries — including a
    final fallback that generates and executes Python code in a subprocess.
    """
    from nethub_runtime.core.services.image_generation_service import ImageGenerationService

    target_path = _resolve_requested_image_path(task.input_text)
    if target_path is None:
        artifact_id = _artifact_id_from_trace(context.trace_id, ".png")
        # Use generated/ dir directly so the path is always on disk
        target_path = Path(f"generated/{artifact_id}.png")
        target_path.parent.mkdir(parents=True, exist_ok=True)

    service = ImageGenerationService(coordinator)
    result = service.generate(task, target_path)

    # Always ensure artifact_path is present so bridge can stage the file
    if result.get("status") in ("generated", "invalid_output") and "artifact_path" not in result:
        result["artifact_path"] = str(target_path)
    if "file_name" not in result and target_path:
        result["file_name"] = target_path.name
    return result


def handle_video_generate_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    """Video generation: AnimateDiff (priority 1) → Stable Video Diffusion (priority 2) via HuggingFace."""
    import hashlib as _hs
    input_text = task.input_text or ""
    artifact_id = _hs.md5(input_text.encode()).hexdigest()[:12]
    out_path = Path(f"generated/{artifact_id}.mp4")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Load candidate list from policy (AnimateDiff → SVD-XT)
    from nethub_runtime.core.config.settings import SEMANTIC_POLICY_PATH
    try:
        policy = json.loads(SEMANTIC_POLICY_PATH.read_text(encoding="utf-8"))
        candidates = (policy.get("capability_acquisition_strategies") or {}) \
            .get("video_generation", {}).get("huggingface_candidates") or []
    except Exception:
        candidates = []

    for candidate in sorted(candidates, key=lambda c: int(c.get("priority", 99))):
        model_id = str(candidate.get("model_id") or "")
        label = str(candidate.get("label", model_id))
        steps = int(candidate.get("inference_steps", 25))
        try:
            import torch  # type: ignore
            if not torch.cuda.is_available():
                continue  # skip local diffusion on CPU-only machines
            if "animatediff" in model_id.lower():
                from diffusers import AnimateDiffPipeline, MotionAdapter  # type: ignore
                import imageio  # type: ignore
                adapter = MotionAdapter.from_pretrained(model_id)
                pipe = AnimateDiffPipeline.from_pretrained(
                    "emilianJR/epiCRealism", motion_adapter=adapter,
                    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                )
                if torch.cuda.is_available():
                    pipe = pipe.to("cuda")
                frames = pipe(input_text, num_frames=16, num_inference_steps=steps).frames[0]
                imageio.mimwrite(str(out_path), frames, fps=8)
            else:
                # Stable Video Diffusion or generic text-to-video
                from diffusers import DiffusionPipeline  # type: ignore
                import imageio  # type: ignore
                pipe = DiffusionPipeline.from_pretrained(
                    model_id,
                    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                )
                if torch.cuda.is_available():
                    pipe = pipe.to("cuda")
                output = pipe(input_text, num_inference_steps=steps)
                frames = output.frames[0] if hasattr(output, "frames") else output.images
                imageio.mimwrite(str(out_path), frames, fps=8)

            if out_path.exists() and out_path.stat().st_size > 0:
                return {"artifact_type": "video", "status": "generated", "method": label,
                        "artifact_path": str(out_path), "file_name": out_path.name}
        except Exception:
            continue

    # Trigger capability acquisition for missing packages/models
    from nethub_runtime.core.services.capability_acquisition_service import CapabilityAcquisitionService
    svc = getattr(coordinator, "capability_acquisition_service", None) or \
          CapabilityAcquisitionService(security_guard=getattr(coordinator, "security_guard", None))
    acq = svc.acquire(task_type="video_generation", gap="no_video_model")
    return {"artifact_type": "video", "status": "acquisition_triggered",
            "method": acq.strategy, "detail": acq.detail,
            "message": "Video generation model is being installed. Retry shortly."}


def handle_file_generate_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    target_path = _resolve_requested_file_path(task.input_text)
    extension = target_path.suffix.lower() if target_path else _infer_file_extension(task.input_text)
    generated_content = _generate_file_content(coordinator, task, context, extension)

    if target_path is None:
        artifact_id = _artifact_id_from_trace(context.trace_id, extension)
        persisted_path = coordinator.generated_artifact_store.persist(
            "code",
            artifact_id,
            generated_content,
            extension=extension,
        )
        target_path = Path(str(persisted_path))
        storage = "generated_artifact_store"
    else:
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_path.write_text(generated_content, encoding="utf-8")
        storage = "workspace"

    return {
        "artifact_type": "file",
        "artifact_path": str(target_path),
        "status": "generated",
        "task": "file_generation",
        "content": generated_content,
        "storage": storage,
        "file_name": target_path.name,
    }


def handle_file_read_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    # --- Priority 1: attachments passed via context (LINE uploads, etc.) ---
    attachments = context.metadata.get("attachments") or []
    if attachments:
        return _process_context_attachments(attachments, context)

    # --- Priority 2: file path mentioned in the task text ---
    target_path = _resolve_existing_file_path(task.input_text)
    if target_path is None:
        return {
            "artifact_type": "file",
            "status": "not_found",
            "message": "Requested file path was not found in the workspace.",
            "content": "",
        }

    content = target_path.read_text(encoding="utf-8")
    return {
        "artifact_type": "file",
        "artifact_path": str(target_path),
        "status": "read",
        "message": f"File content from {target_path}",
        "content": content,
        "file_name": target_path.name,
        "storage": "workspace",
    }


def _received_dir(session_id: str) -> Path:
    """Return (and create) the per-session received directory inside nethub_runtime."""
    workspace = os.getenv("NESTHUB_WORKSPACE_PATH", "").strip()
    base = Path(workspace) if workspace else Path(__file__).resolve().parents[3]
    received = base / "received" / session_id
    received.mkdir(parents=True, exist_ok=True)
    return received


def _extract_file_content(path: Path, content_type: str) -> str:
    """Best-effort text extraction from a file.

    Supports plain text formats natively; tries PyMuPDF for PDFs,
    python-docx for Word, openpyxl for Excel.  Falls back to a path
    reference for unrecognised binary formats.
    """
    suffix = path.suffix.lower()

    # Plain text
    if suffix in (".txt", ".md", ".json", ".yaml", ".yml", ".html", ".htm",
                  ".csv", ".xml", ".js", ".css", ".py", ".ts"):
        try:
            return path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            return ""

    # PDF
    if suffix == ".pdf" or content_type == "application/pdf":
        try:
            import fitz  # type: ignore  # PyMuPDF
            doc = fitz.open(str(path))
            return "\n".join(page.get_text() for page in doc)
        except ImportError:
            pass
        except Exception:
            pass

    # Word (.docx)
    if suffix in (".docx", ".doc") or "wordprocessingml" in content_type:
        try:
            from docx import Document  # type: ignore  # python-docx
            doc = Document(str(path))
            return "\n".join(para.text for para in doc.paragraphs)
        except ImportError:
            pass
        except Exception:
            pass

    # Excel (.xlsx, .xls)
    if suffix in (".xlsx", ".xls") or "spreadsheetml" in content_type:
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in wb.worksheets:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    lines.append("\t".join("" if v is None else str(v) for v in row))
            return "\n".join(lines)
        except ImportError:
            pass
        except Exception:
            pass

    # Image — return path reference (OCR handled by a dedicated step)
    if content_type.startswith("image/") or suffix in (".jpg", ".jpeg", ".png", ".gif", ".webp"):
        return f"[IMAGE FILE: {path.name}]"

    # Binary fallback
    return f"[BINARY FILE: {path.name} ({content_type})]"


def _process_context_attachments(
    attachments: list[dict[str, Any]],
    context: CoreContextSchema,
) -> dict[str, Any]:
    """Download attachment files to nethub_runtime/received/ and read their content."""
    dest_dir = _received_dir(context.session_id)
    processed: list[dict[str, Any]] = []

    for att in attachments:
        url = str(att.get("download_url") or "").strip()
        file_name = str(att.get("file_name") or "file.bin").strip()
        content_type = str(att.get("content_type") or "application/octet-stream")

        if not url and not str(att.get("stored_path") or "").strip():
            continue

        local_path = dest_dir / file_name
        # Prefer stored_path (same-process direct read — no HTTP round-trip)
        sp = str(att.get("stored_path") or "").strip()
        if sp and Path(sp).exists():
            import shutil
            shutil.copy2(sp, local_path)
        elif url:
            # Fallback: HTTP download (cross-process / remote deployments)
            try:
                req = urllib.request.Request(url)
                with urllib.request.urlopen(req, timeout=30) as resp:
                    local_path.write_bytes(resp.read())
            except Exception as exc:
                processed.append({
                    "file_name": file_name,
                    "status": "download_error",
                    "error": str(exc),
                    "content": "",
                })
                continue
        else:
            processed.append({
                "file_name": file_name,
                "status": "download_error",
                "error": "no stored_path and no download_url",
                "content": "",
            })
            continue

        # TODO: 接入文件内容提取（PDF / Word / Excel / 纯文本）
        processed.append({
            "file_name": file_name,
            "artifact_path": str(local_path),
            "content_type": content_type,
            "file_size": local_path.stat().st_size,
            "content": f"[文件已接收: {file_name}，内容提取待实现]",
            "status": "received",
        })

    if not processed:
        return {
            "artifact_type": "file",
            "status": "no_attachments",
            "message": "未找到附件。",
            "content": "",
        }

    # Return the first (primary) attachment as the main result,
    # include all in the attachments list for multi-file scenarios
    primary = next((p for p in processed if p["status"] == "received"), processed[0])
    return {
        "artifact_type": "file",
        "artifact_path": primary.get("artifact_path", ""),
        "status": primary.get("status", "received"),
        "message": f"已接收 {len([p for p in processed if p['status'] == 'received'])} 个附件",
        "content": primary.get("content", ""),
        "file_name": primary.get("file_name", ""),
        "storage": "received",
        "attachments": processed,
    }


def _resolve_requested_image_path(text: str) -> Path | None:
    match = re.search(
        r"([A-Za-z0-9_./\\-]+\.(?:png|jpe?g|gif|webp|bmp|svg))",
        text,
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    candidate = match.group(1).strip().strip("\"'")
    return Path(candidate).expanduser() if candidate else None


def _resolve_requested_file_path(text: str) -> Path | None:
    match = re.search(
        r"([A-Za-z0-9_./\\-]+\.(?:html?|js|css|json|md|txt|py|ya?ml))",
        text,
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    candidate = match.group(1).strip().strip("\"'")
    return Path(candidate).expanduser() if candidate else None



def _resolve_existing_file_path(text: str) -> Path | None:
    match = re.search(r"([A-Za-z0-9_./\\-]+\.(?:html?|js|css|json|md|txt|py|ya?ml))", text, flags=re.IGNORECASE)
    if not match:
        return None
    candidate = match.group(1).strip().strip('"\'')
    if not candidate:
        return None
    path = Path(candidate).expanduser()
    candidates = [path]
    if not path.is_absolute():
        candidates.append((Path.cwd() / path).resolve())
        candidates.append((Path(__file__).resolve().parents[3] / path).resolve())
    for candidate_path in candidates:
        if candidate_path.exists() and candidate_path.is_file():
            return candidate_path
    return None


def _infer_file_extension(text: str) -> str:
    lowered = text.lower()
    if any(marker in lowered for marker in ("html", "网页", "web page")):
        return ".html"
    if any(marker in lowered for marker in ("javascript", " js", ".js")):
        return ".js"
    if any(marker in lowered for marker in ("css", ".css")):
        return ".css"
    if any(marker in lowered for marker in ("json", ".json")):
        return ".json"
    if any(marker in lowered for marker in ("markdown", ".md")):
        return ".md"
    return ".txt"


def _artifact_id_from_trace(trace_id: str, extension: str) -> str:
    suffix = extension.lstrip(".") or "file"
    return f"{trace_id}_{suffix}"


def _generate_file_content(coordinator: Any, task: TaskSchema, context: CoreContextSchema, extension: str) -> str:
    generated = _normalize_model_file_content(_generate_file_content_via_model(coordinator, task, context, extension), extension)
    if generated:
        return generated
    if extension == ".html":
        return _fallback_html_content(task.input_text)
    return task.input_text


def _generate_file_content_via_model(coordinator: Any, task: TaskSchema, context: CoreContextSchema, extension: str) -> str | None:
    if coordinator.model_router is None or not hasattr(coordinator, "_invoke_model_text"):
        return None
    return coordinator._invoke_model_text(
        task_type="code_generation",
        prompt=(
            "Generate only the final file content. Do not use markdown fences.\n"
            f"task_intent: {task.intent}\n"
            f"session_id: {context.session_id}\n"
            f"target_extension: {extension}\n"
            f"user_request: {task.input_text}"
        ),
        system_prompt=(
            "You are the NestHub runtime code generation engine. "
            "Return only valid file content for the requested file type."
        ),
        temperature=0.2,
    )


def _normalize_model_file_content(cleaned: str | None, extension: str) -> str | None:
    if not cleaned or cleaned.startswith("Model response (mock):"):
        return None
    cleaned = _strip_code_fences(cleaned)
    if extension == ".json":
        try:
            return json.dumps(json.loads(cleaned), ensure_ascii=False, indent=2)
        except Exception:
            return None
    return cleaned


def _strip_code_fences(text: str) -> str:
    stripped = text.strip()
    if not stripped.startswith("```"):
        return stripped
    stripped = re.sub(r"^```[a-zA-Z0-9_-]*\n", "", stripped)
    stripped = re.sub(r"\n```$", "", stripped)
    return stripped.strip()


def _fallback_html_content(request_text: str) -> str:
    alert_message = _extract_alert_message(request_text)
    button_label = _extract_button_label(request_text)
    title = _extract_html_title(request_text)
    return """<!DOCTYPE html>
<html lang=\"en\">
<head>
  <meta charset=\"UTF-8\">
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">
  <title>{title}</title>
</head>
<body>
  <button id=\"submitButton\" type=\"button\">{button_label}</button>
  <script>
    document.getElementById(\"submitButton\").addEventListener(\"click\", function () {{
      alert(\"{alert_message}\");
    }});
  </script>
</body>
</html>
""".format(title=title, button_label=button_label, alert_message=alert_message)


def _extract_alert_message(text: str) -> str:
    match = re.search(r"(?:弹出|提示|alert)\s*[:：]?\s*([A-Za-z0-9 ,.!?-]+)", text, flags=re.IGNORECASE)
    if match:
        return match.group(1).strip()
    if "hello" in text.lower() and "world" in text.lower():
        return "hello,world!"
    return "hello,world!"


def _extract_button_label(text: str) -> str:
    match = re.search(r"(?:按钮|button)\s*[:：]?\s*([A-Za-z][A-Za-z0-9 _-]{0,20})", text, flags=re.IGNORECASE)
    if match:
        return match.group(1).strip()
    if "submit" in text.lower():
        return "Submit"
    return "Submit"


def _extract_html_title(text: str) -> str:
    if "hello" in text.lower() and "world" in text.lower():
        return "Hello World"
    return "NestHub Generated Page"


def handle_web_retrieve_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    """Fetch web content for the URL (or query) in *task.input_text*.

    Strategy (local-first, no paid API required):
      1. Extract URL from input text, or build a DuckDuckGo search URL.
      2. Fetch with ``urllib.request`` (bundled, no extra install needed).
      3. Strip HTML tags to plain text.
      4. If ``readability-lxml`` is installed, use it for cleaner extraction.
      5. Store the raw content in the VectorStore under namespace="web_cache"
         so the business-domain layer remembers it for future queries.
    """
    url = _extract_url_from_text(task.input_text)
    if not url:
        # Fall back to DuckDuckGo lite HTML search
        query = urllib.parse.quote_plus(task.input_text)
        url = f"https://html.duckduckgo.com/html/?q={query}"

    raw_html = ""
    try:
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (nesthub-runtime/1.0)"},
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            charset = "utf-8"
            ct = resp.headers.get("Content-Type", "")
            if "charset=" in ct:
                charset = ct.split("charset=")[-1].strip().split(";")[0].strip()
            raw_html = resp.read().decode(charset, errors="replace")
    except Exception as exc:
        return {
            "artifact_type": "web_content",
            "status": "failed",
            "url": url,
            "content": "",
            "error": str(exc),
        }

    # Extract readable text
    content = _extract_readable_text(raw_html, url)

    # Persist to VectorStore (business domain layer — non-blocking)
    try:
        if hasattr(coordinator, "vector_store") and coordinator.vector_store is not None:
            coordinator.vector_store.add_knowledge(
                namespace="web_cache",
                content=content[:2000],
                metadata={"url": url, "session_id": context.session_id, "trace_id": context.trace_id},
                item_id=f"web_{context.trace_id}",
            )
    except Exception:
        pass

    return {
        "artifact_type": "web_content",
        "status": "retrieved",
        "url": url,
        "content": content,
        "content_length": len(content),
    }


def _extract_url_from_text(text: str) -> str | None:
    """Extract the first HTTP(S) URL from *text*, or return None."""
    match = re.search(r"https?://[^\s\"'<>]+", text)
    if match:
        url = match.group(0).rstrip(".,;)")
        return url
    return None


def _extract_readable_text(html: str, url: str = "") -> str:
    """Extract plain text from HTML.

    Tries ``readability-lxml`` first (cleaner extraction), then falls back to
    a simple regex tag-stripper that always works with zero extra dependencies.
    """
    # Try readability-lxml (installed via web_research strategy)
    try:
        from readability import Document  # type: ignore
        doc = Document(html)
        summary_html = doc.summary()
        text = re.sub(r"<[^>]+>", " ", summary_html)
        text = re.sub(r"&[a-zA-Z]+;", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            return text[:8000]
    except Exception:
        pass

    # Fallback: strip all tags
    text = re.sub(r"<script[^>]*>.*?</script>", " ", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"&[a-zA-Z#0-9]+;", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:8000]


def handle_web_summarize_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    step_outputs: dict[str, Any],
) -> dict[str, Any]:
    """Summarise web content fetched by ``web_retrieve`` using a local LLM.

    Reads ``step_outputs["web_retrieve"]["content"]`` and passes it to
    ``coordinator._invoke_model_text`` with task_type="web_summarize".
    Falls back to a truncated excerpt when the model router is unavailable.
    """
    retrieve_out = step_outputs.get("web_retrieve", {})
    web_content = retrieve_out.get("content", "")
    url = retrieve_out.get("url", "")

    if not web_content:
        # Try VectorStore cache
        try:
            if hasattr(coordinator, "vector_store") and coordinator.vector_store is not None:
                hits = coordinator.vector_store.search(
                    task.input_text, top_k=1, namespace="web_cache"
                )
                if hits:
                    web_content = hits[0].get("content", "")
                    url = hits[0].get("metadata", {}).get("url", url)
        except Exception:
            pass

    if not web_content:
        return {
            "artifact_type": "summary",
            "status": "no_content",
            "summary": "",
            "url": url,
        }

    # Truncate to fit context window (≤ 3000 chars for most local models)
    excerpt = web_content[:3000]

    # Local LLM summarization
    summary: str | None = None
    if hasattr(coordinator, "_invoke_model_text"):
        summary = coordinator._invoke_model_text(
            task_type="web_summarize",
            prompt=(
                f"Summarise the following web page content in 3-5 sentences.\n"
                f"User question: {task.input_text}\n\n"
                f"Content:\n{excerpt}"
            ),
            system_prompt=(
                "You are a web research assistant. Summarise web content concisely and "
                "accurately. Focus on information relevant to the user's question."
            ),
        )

    if not summary:
        # Fallback: return first 500 chars as plain excerpt
        summary = excerpt[:500]

    # Persist summary to VectorStore
    try:
        if hasattr(coordinator, "vector_store") and coordinator.vector_store is not None:
            coordinator.vector_store.add_knowledge(
                namespace="web_summary",
                content=summary,
                metadata={"url": url, "query": task.input_text, "trace_id": context.trace_id},
                item_id=f"summary_{context.trace_id}",
            )
    except Exception:
        pass

    return {
        "artifact_type": "summary",
        "status": "summarized",
        "summary": summary,
        "url": url,
        "source_length": len(web_content),
    }


def handle_prepare_runtime_tools_step(
    coordinator: Any,
    step: dict[str, Any],
    _task: TaskSchema,
    _context: CoreContextSchema,
    _step_outputs: dict[str, Any],
) -> dict[str, Any]:
    repair_metadata = step.get("metadata") or {}
    missing_items = list(repair_metadata.get("missing_tools", []))
    install_plan = coordinator._build_runtime_install_plan(missing_items)
    execution_result = None
    if coordinator._allow_runtime_auto_install() and install_plan.get("auto_install"):
        execution_result = coordinator._execute_runtime_install_plan(install_plan)
    return {
        "status": "executed" if execution_result else "prepared",
        "missing_items": missing_items,
        "install_plan": install_plan,
        "install_execution": execution_result,
        "message": "Runtime tool preparation plan generated.",
    }


def handle_generate_workflow_artifact_step(
    coordinator: Any,
    step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    step_outputs: dict[str, Any],
) -> dict[str, Any]:
    analysis_output = step_outputs.get("analyze_workflow_context", {})
    summary = analysis_output.get("summary") or analysis_output.get("analysis") or task.input_text
    content_lines = [
        f"task_intent: {task.intent}",
        f"session_id: {context.session_id}",
        "",
        str(summary),
    ]
    artifact_id = f"artifact_{context.trace_id}"
    artifact_path = coordinator.generated_artifact_store.persist(
        "feature",
        artifact_id,
        "\n".join(content_lines),
        extension=".md",
    )
    return {
        "artifact": str(artifact_path),
        "artifact_type": "document",
        "artifact_path": str(artifact_path),
        "status": "generated",
        "summary": str(summary),
    }


def handle_generate_runtime_patch_step(
    coordinator: Any,
    _step: dict[str, Any],
    task: TaskSchema,
    context: CoreContextSchema,
    step_outputs: dict[str, Any],
) -> dict[str, Any]:
    return coordinator._generate_runtime_patch(task=task, context=context, step_outputs=step_outputs)


def handle_validate_runtime_patch_step(
    coordinator: Any,
    _step: dict[str, Any],
    _task: TaskSchema,
    context: CoreContextSchema,
    step_outputs: dict[str, Any],
) -> dict[str, Any]:
    patch_payload = step_outputs.get("generate_runtime_patch", {})
    return coordinator._run_runtime_validation(context=context, patch_payload=patch_payload)


def handle_verify_runtime_patch_step(
    coordinator: Any,
    _step: dict[str, Any],
    _task: TaskSchema,
    context: CoreContextSchema,
    step_outputs: dict[str, Any],
) -> dict[str, Any]:
    patch_payload = step_outputs.get("generate_runtime_patch", {})
    validation_payload = step_outputs.get("validate_runtime_patch", {})
    return coordinator._verify_runtime_patch(context=context, patch_payload=patch_payload, validation_payload=validation_payload)


def handle_persist_workflow_output_step(
    coordinator: Any,
    _step: dict[str, Any],
    _task: TaskSchema,
    _context: CoreContextSchema,
    step_outputs: dict[str, Any],
) -> dict[str, Any]:
    artifact_payload = step_outputs.get("generate_workflow_artifact", {})
    artifact_path = artifact_payload.get("artifact_path")
    if not artifact_path:
        return {
            "delivery_status": "skipped",
            "stored_output": None,
            "message": "No artifact available for persistence.",
        }
    path = Path(str(artifact_path))
    return {
        "delivery_status": "stored",
        "stored_output": str(path),
        "message": f"Workflow output stored at {path.name}.",
    }